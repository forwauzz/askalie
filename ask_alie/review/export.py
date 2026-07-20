"""Chronology exports: JSON, CSV and a self-contained HTML file (Spec §27)."""

from __future__ import annotations

import csv
import html
import json
from typing import Any

from ask_alie.review.service import build_rows, split_queues
from ask_alie.workspace.paths import CasePaths

_CSV_COLUMNS = ["Date", "Type", "Description", "Auteur", "Source", "Page", "File", "Statut"]

_MONTHS_FR = [
    "janvier", "février", "mars", "avril", "mai", "juin", "juillet",
    "août", "septembre", "octobre", "novembre", "décembre",
]


def _date_fr(iso: str | None) -> str:
    if not iso:
        return "Date non résolue"
    try:
        from datetime import date as _date

        d = _date.fromisoformat(iso)
        return f"{d.day}{'er' if d.day == 1 else ''} {_MONTHS_FR[d.month - 1]} {d.year}"
    except ValueError:
        return iso


def export_all(paths: CasePaths) -> dict[str, str]:
    rows = build_rows(paths)
    queues = split_queues(rows)
    paths.output_dir.mkdir(parents=True, exist_ok=True)

    json_path = paths.output_dir / "chronology.json"
    json_path.write_text(json.dumps(rows, ensure_ascii=False, indent=2), encoding="utf-8")

    csv_path = paths.output_dir / "chronology.csv"
    with csv_path.open("w", encoding="utf-8-sig", newline="") as fh:
        writer = csv.writer(fh)
        writer.writerow(_CSV_COLUMNS)
        for row in queues["default"] + queues["secondary"]:
            writer.writerow(
                [
                    row["date"] or "date non résolue",
                    row["event_type"],
                    row["summary"],
                    row["author"] or "",
                    row["report_id"],
                    ", ".join(str(p) for p in row["source_pages"]),
                    row["queue"],
                    row["status"],
                ]
            )

    html_path = paths.output_dir / "chronology.html"
    html_path.write_text(_render_html(queues), encoding="utf-8")

    docx_path = export_docx(paths, rows=rows, queues=queues)
    return {
        "json": str(json_path),
        "csv": str(csv_path),
        "html": str(html_path),
        "docx": docx_path,
    }


def export_docx(
    paths: CasePaths,
    rows: list[dict[str, Any]] | None = None,
    queues: dict[str, list[dict[str, Any]]] | None = None,
) -> str:
    """Clean Word chronology: title block, then one table per queue (French)."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    if queues is None:
        rows = rows if rows is not None else build_rows(paths)
        queues = split_queues(rows)

    doc = Document()
    for section in doc.sections:
        section.left_margin = section.right_margin = Cm(2)
        section.top_margin = section.bottom_margin = Cm(2)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    run = title.add_run("Chronologie médico-légale")
    run.bold = True
    run.font.size = Pt(20)
    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run(
        f"Dossier : {paths.root.name} — {len(queues['default'])} événements principaux, "
        f"{len(queues['secondary'])} secondaires, {len(queues['unresolved'])} à réviser. "
        "Générée par Ask ALIE; chaque événement cite sa source."
    )
    sub_run.font.size = Pt(9.5)
    sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x6E)

    def add_table(heading: str, table_rows: list[dict[str, Any]], note: str = "") -> None:
        doc.add_heading(f"{heading} ({len(table_rows)})", level=1)
        if note:
            note_par = doc.add_paragraph()
            note_run = note_par.add_run(note)
            note_run.italic = True
            note_run.font.size = Pt(9)
        if not table_rows:
            doc.add_paragraph("Aucun événement.")
            return
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        widths = (Cm(3.4), Cm(9.6), Cm(4.0))
        for cell, text in zip(table.rows[0].cells, ("Date", "Événement", "Source")):
            cell.paragraphs[0].add_run(text).bold = True
        for row in table_rows:
            cells = table.add_row().cells
            date_par = cells[0].paragraphs[0]
            date_par.add_run(_date_fr(row["date"])).bold = True

            event_par = cells[1].paragraphs[0]
            event_par.add_run(row["summary"] or "—")
            meta_bits = [b for b in (row["event_type"], row["author"]) if b]
            if row["needs_review"]:
                meta_bits.append("à réviser")
            if meta_bits:
                meta_par = cells[1].add_paragraph()
                meta_run = meta_par.add_run(" · ".join(meta_bits))
                meta_run.font.size = Pt(8.5)
                meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x6E)
            if row["quote"]:
                quote_par = cells[1].add_paragraph()
                quote_run = quote_par.add_run(f"« {row['quote'].strip()} »")
                quote_run.italic = True
                quote_run.font.size = Pt(9)
                quote_run.font.color.rgb = RGBColor(0x4C, 0x4C, 0x55)

            pages = ", ".join(str(p) for p in row["source_pages"]) or "—"
            source_par = cells[2].paragraphs[0]
            source_par.add_run(row["source_document"]).font.size = Pt(9)
            page_par = cells[2].add_paragraph()
            page_run = page_par.add_run(f"p. {pages}")
            page_run.font.size = Pt(8.5)
            page_run.font.color.rgb = RGBColor(0x66, 0x66, 0x6E)
            for cell, width in zip(cells, widths):
                cell.width = width
        for cell, width in zip(table.rows[0].cells, widths):
            cell.width = width

    add_table("File principale", queues["default"])
    doc.add_page_break()
    add_table("File secondaire", queues["secondary"],
              "Événements documentés mais de moindre priorité; aucun n'est supprimé.")
    doc.add_page_break()
    add_table("Non résolu / à réviser", queues["unresolved"],
              "Dates non résolues, doublons potentiels et événements nécessitant une validation.")

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Document de travail — vérifier chaque entrée à la source avant usage juridique.")
    footer_run.font.size = Pt(8.5)
    footer_run.italic = True

    paths.output_dir.mkdir(parents=True, exist_ok=True)
    docx_path = paths.output_dir / "chronology.docx"
    doc.save(docx_path)
    return str(docx_path)


def _table(rows: list[dict[str, Any]]) -> str:
    if not rows:
        return "<p>Aucun événement.</p>"
    cells = []
    for row in rows:
        quote = html.escape(row["quote"] or "")
        pages = ", ".join(str(p) for p in row["source_pages"])
        cells.append(
            "<tr>"
            f"<td>{html.escape(row['date'] or 'date non résolue')}</td>"
            f"<td>{html.escape(row['event_type'])}</td>"
            f"<td>{html.escape(row['summary'])}"
            + (
                f"<details><summary>citation (p. {row['quote_page']})</summary>"
                f"<blockquote>{quote}</blockquote></details>"
                if quote
                else ""
            )
            + "</td>"
            f"<td>{html.escape(row['author'] or '')}</td>"
            f"<td>{html.escape(row['report_id'])} p. {pages}</td>"
            f"<td>{html.escape(row['status'])}</td>"
            "</tr>"
        )
    return (
        "<table><thead><tr><th>Date</th><th>Type</th><th>Description</th>"
        "<th>Auteur</th><th>Source</th><th>Statut</th></tr></thead><tbody>"
        + "".join(cells)
        + "</tbody></table>"
    )


def _render_html(queues: dict[str, list[dict[str, Any]]]) -> str:
    return f"""<!doctype html>
<html lang="fr"><head><meta charset="utf-8"><title>Chronologie Ask ALIE</title>
<style>
body {{ font-family: system-ui, sans-serif; margin: 2rem; color: #1a1a1a; }}
table {{ border-collapse: collapse; width: 100%; margin-bottom: 2rem; }}
th, td {{ border: 1px solid #ccc; padding: 6px 8px; text-align: left; vertical-align: top; }}
th {{ background: #f0f0f0; }}
blockquote {{ margin: 4px 0; padding-left: 8px; border-left: 3px solid #999; color: #444; }}
h2 {{ margin-top: 2rem; }}
</style></head><body>
<h1>Chronologie — Ask ALIE (POC)</h1>
<h2>File principale ({len(queues["default"])})</h2>
{_table(queues["default"])}
<h2>File secondaire ({len(queues["secondary"])})</h2>
{_table(queues["secondary"])}
<h2>Non résolu / à réviser ({len(queues["unresolved"])})</h2>
{_table(queues["unresolved"])}
</body></html>
"""
